from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

output = Path(r"c:\xampp\php\www\keuangan-app\docs\pengujian-nlp-untuk-ppt.docx")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Arial"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
style.font.size = Pt(11)


def add_center_title(text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_heading(text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


add_center_title("Pengujian NLP pada Aplikasi E-Keuangan")

doc.add_paragraph(
    "Dokumen ini disusun sebagai bahan penjelasan untuk Word dan nantinya dapat dipindahkan ke PowerPoint. "
    "Fokusnya adalah menjelaskan alur pengujian pada halaman utama aplikasi E-Keuangan MAN 2 Surakarta, "
    "mulai dari login sampai modul operasional yang paling sering digunakan."
)

add_heading("1. Tujuan Pengujian", 1)
doc.add_paragraph(
    "Pengujian dilakukan untuk memastikan setiap halaman utama aplikasi berjalan sesuai fungsi, mudah digunakan, "
    "dan menampilkan data yang benar. Dari sisi presentasi, isi dokumen ini bisa dipakai untuk menjelaskan bahwa "
    "sistem sudah melewati pengujian dasar pada bagian autentikasi, dashboard, manajemen siswa, master bayar, "
    "dan manajemen tagihan."
)

add_heading("2. Ruang Lingkup Pengujian", 1)
add_bullets([
    "Login pengguna.",
    "Dashboard ringkasan keuangan.",
    "Manajemen data siswa.",
    "Master bayar.",
    "Manajemen tagihan.",
])

add_heading("3. Deskripsi Hasil Pengujian per Modul", 1)

modules = [
    (
        "3.1 Login",
        "Halaman login merupakan pintu masuk utama ke sistem. Pada halaman ini pengguna memasukkan username atau email dan kata sandi untuk mengakses aplikasi. Tampilan login sudah dibuat jelas, sederhana, dan fokus pada proses autentikasi.",
        [
            "Login digunakan sebagai validasi awal agar hanya pengguna yang memiliki akun yang bisa masuk.",
            "Halaman login menampilkan identitas aplikasi dan institusi sehingga terlihat resmi dan mudah dikenali.",
            "Form login dibuat ringkas agar proses masuk ke sistem lebih cepat.",
        ],
        "Pada tahap awal, pengguna diarahkan ke halaman login. Di sini pengguna harus memasukkan username atau email beserta kata sandi. Jika data yang dimasukkan sesuai, sistem akan mengarahkan pengguna ke dashboard utama.",
    ),
    (
        "3.2 Dashboard",
        "Dashboard berfungsi sebagai pusat informasi utama setelah pengguna berhasil login. Pada halaman ini ditampilkan ringkasan keuangan seperti total pembayaran, sisa tagihan, saldo kas bersih, tren pembayaran, dan daftar tunggakan utama.",
        [
            "Dashboard membantu pengguna melihat kondisi keuangan secara cepat tanpa harus membuka laporan satu per satu.",
            "Informasi yang ditampilkan bersifat ringkas namun penting untuk pengambilan keputusan.",
            "Grafik dan kartu ringkasan memudahkan pengguna memahami kondisi transaksi dalam satu tampilan.",
        ],
        "Dashboard menjadi pusat kontrol sistem. Dari halaman ini, admin dapat langsung memantau total pembayaran, sisa tagihan, saldo kas, serta perkembangan pembayaran selama beberapa bulan terakhir.",
    ),
    (
        "3.3 Manajemen Siswa",
        "Menu manajemen siswa digunakan untuk mengelola data siswa yang aktif maupun nonaktif. Pada halaman ini tersedia fitur pencarian, filter berdasarkan angkatan, kelas, jurusan, dan tipe siswa, serta aksi untuk menambah, mengubah, mengaktifkan, atau menonaktifkan data siswa.",
        [
            "Modul ini mempermudah administrasi data siswa dalam satu halaman.",
            "Data siswa dapat difilter agar pencarian lebih cepat dan lebih tepat.",
            "Status siswa aktif atau nonaktif dapat dipantau dengan jelas.",
        ],
        "Pada menu manajemen siswa, admin dapat melihat daftar siswa secara lengkap sekaligus melakukan filter berdasarkan angkatan, kelas, jurusan, dan tipe siswa. Hal ini membuat pengelolaan data menjadi lebih terstruktur dan efisien.",
    ),
    (
        "3.4 Master Bayar",
        "Master bayar digunakan untuk mengatur komponen biaya dan skema tarif. Di dalamnya terdapat pengelolaan jenis biaya, seperti SPP atau biaya lain, serta pengaturan nominal dan periode berlaku untuk tiap angkatan atau kategori tertentu.",
        [
            "Master bayar adalah dasar pembentukan tagihan siswa.",
            "Admin dapat membuat jenis biaya baru sesuai kebutuhan sekolah.",
            "Skema tarif bisa disesuaikan berdasarkan angkatan, jenis biaya, dan masa berlaku.",
        ],
        "Menu master bayar dipakai untuk mendefinisikan struktur biaya sekolah. Dari sini admin menentukan jenis biaya apa saja yang akan ditagihkan, besarnya nominal, dan untuk siapa biaya tersebut berlaku.",
    ),
    (
        "3.5 Manajemen Tagihan",
        "Manajemen tagihan digunakan untuk membuat siklus tagihan dan menghasilkan invoice secara massal. Modul ini menjadi penghubung antara master biaya dengan tagihan yang diterima siswa.",
        [
            "Admin dapat membuat siklus tagihan per periode tertentu.",
            "Invoice dapat digenerate berdasarkan jenis biaya dan filter siswa.",
            "Modul ini membantu proses penagihan menjadi lebih cepat dan terkontrol.",
        ],
        "Pada menu manajemen tagihan, admin membuat siklus tagihan terlebih dahulu, kemudian sistem menghasilkan invoice sesuai jenis biaya dan filter yang dipilih. Dengan cara ini, proses penagihan siswa menjadi lebih sistematis.",
    ),
]

for title, body, bullets, quote in modules:
    add_heading(title, 2)
    doc.add_paragraph(body)
    p = doc.add_paragraph()
    run = p.add_run("Poin penting untuk PPT:")
    run.bold = True
    add_bullets(bullets)
    p = doc.add_paragraph()
    run = p.add_run("Contoh penjelasan saat presentasi:")
    run.bold = True
    doc.add_paragraph(quote)

add_heading("4. Ringkasan Alur Sistem", 1)
doc.add_paragraph("Alur penggunaan aplikasi dapat dijelaskan seperti berikut:")
add_numbers([
    "Pengguna melakukan login.",
    "Sistem menampilkan dashboard ringkasan keuangan.",
    "Admin mengelola data siswa jika ada perubahan data atau penambahan siswa baru.",
    "Admin mengatur jenis biaya dan skema tarif melalui master bayar.",
    "Admin membuat siklus dan invoice pada menu manajemen tagihan.",
])
doc.add_paragraph(
    "Alur ini menunjukkan bahwa aplikasi sudah memiliki struktur kerja yang jelas, mulai dari data dasar, pengaturan biaya, hingga pembentukan tagihan."
)

add_heading("5. Kalimat Siap Pakai untuk PPT", 1)
add_bullets([
    "Halaman login berfungsi sebagai gerbang utama untuk menjaga keamanan sistem.",
    "Dashboard menampilkan ringkasan kondisi keuangan dalam bentuk yang mudah dibaca.",
    "Manajemen siswa memudahkan admin mengelola data siswa aktif maupun nonaktif.",
    "Master bayar digunakan untuk menyusun komponen biaya dan skema tarif sekolah.",
    "Manajemen tagihan menghubungkan data biaya dengan invoice yang diterima siswa.",
])

add_heading("6. Penutup", 1)
doc.add_paragraph(
    "Secara keseluruhan, hasil pengujian menunjukkan bahwa modul utama aplikasi E-Keuangan sudah memiliki alur kerja yang jelas dan saling terhubung. Login menjadi pintu masuk, dashboard menjadi pusat informasi, manajemen siswa menjaga data tetap rapi, master bayar mengatur struktur biaya, dan manajemen tagihan menangani pembentukan invoice."
)
doc.add_paragraph(
    "Dokumen ini bisa langsung dipindahkan ke Word lalu dipecah menjadi beberapa slide PowerPoint sesuai kebutuhan presentasi."
)

doc.save(output)
print(output)
